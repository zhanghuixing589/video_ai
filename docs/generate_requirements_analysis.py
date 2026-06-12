from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("视频流媒体平台需求分析报告.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "EAF4EA"
LIGHT_YELLOW = "FFF4D6"
LIGHT_RED = "FCE8E6"
GRAY = "667085"
BLACK = "202124"
WHITE = "FFFFFF"
FONT = "Microsoft YaHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_dxa = sum(int(width * 1440) for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_dxa = int(widths[index] * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=GRAY)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True)
        body = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(body)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title)
    set_run_font(title_run, size=10.5, bold=True, color=NAVY)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_run = body_paragraph.add_run(body)
    set_run_font(body_run, size=10.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if index == 0 and len(headers) <= 3:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    style_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in style_specs.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[list_style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Video AI Platform  |  需求分析报告")
    set_run_font(run, size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(85)
    paragraph.paragraph_format.space_after = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("视频流媒体平台")
    set_run_font(run, size=30, bold=True, color=NAVY)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(38)
    run = paragraph.add_run("真实需求分析与产品范围说明")
    set_run_font(run, size=17, color=BLUE)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(65)
    run = paragraph.add_run("基于当前代码仓库、运行配置与已实现业务骨架")
    set_run_font(run, size=11, color=GRAY, italic=True)

    add_table(
        doc,
        ["文档属性", "内容"],
        [
            ("项目名称", "Video AI Platform"),
            ("文档类型", "软件需求分析报告（现状审计版）"),
            ("分析基线", "当前工作区代码与 Docker 开发环境"),
            ("分析日期", "2026 年 6 月 11 日"),
            ("适用阶段", "原型完善、MVP 规划、课程答辩、迭代排期"),
        ],
        [1.55, 4.95],
        header_fill=LIGHT_BLUE,
        font_size=10,
    )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(28)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("文档结论：项目方向合理，但当前仍属于业务骨架原型，尚不具备真实流媒体平台的上线条件。")
    set_run_font(run, size=11, bold=True, color=DARK_BLUE)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. 执行摘要", 1)
    add_body(
        doc,
        "本项目定位为面向普通观众、制片厂、审片员和管理员的多角色视频流媒体平台。"
        "平台核心业务不是单纯的视频展示，而是“制片厂准入—片库提交—内容审核—发布播放—运营管理”的闭环，"
        "并计划在后续接入 AI 推荐、语义搜索、字幕生成等能力。"
    )
    add_callout(
        doc,
        "总体判断",
        "当前后端已经具备用户角色、制片厂申请、视频分类和审核状态流转的基础实体与接口；"
        "PostgreSQL、Redis、pgAdmin 和 Redis Commander 的开发环境可运行。"
        "但认证授权仍被完全放行，前端当前无法构建，AI 微服务不存在，视频文件上传、转码、播放和版权治理尚未设计。"
        "因此建议先完成可验证的 MVP，再进入 AI 和大规模流媒体能力建设。",
        LIGHT_YELLOW,
    )

    add_heading(doc, "2. 项目现状审计", 1)
    add_heading(doc, "2.1 当前技术与运行状态", 2)
    add_table(
        doc,
        ["领域", "当前事实", "成熟度"],
        [
            ("Java 后端", "Java 17、Spring Boot 3.2、JPA、Security、Maven；clean compile 已通过。", "基础可用"),
            ("前端", "React、TypeScript、Ant Design、Axios 已安装；主页面文件缺失且 TypeScript 构建失败。", "不可交付"),
            ("数据库", "PostgreSQL 15 容器健康，JPA 可自动建表；使用 ddl-auto=update。", "开发可用"),
            ("Redis", "Redis 7 容器健康，但业务代码尚未实际使用缓存。", "仅基础设施"),
            ("AI 微服务", "README 有规划，仓库中没有 FastAPI/LangChain 服务目录或代码。", "未实现"),
            ("部署", "有 Dockerfile 和 Compose；开发环境可启动数据库管理工具。", "部分实现"),
            ("测试", "只有 Spring 上下文测试骨架，缺少领域、接口和前端测试。", "不足"),
        ],
        [1.15, 4.35, 1.0],
    )

    add_heading(doc, "2.2 已实现业务能力", 2)
    for item in [
        "定义 ADMIN、REVIEWER、STUDIO、USER 四类角色。",
        "支持创建用户、查询用户、修改角色、提交制片厂申请、审核制片厂申请。",
        "视频区分电影、综艺、电视剧，并具有热血、爱情、喜剧、悬疑、科幻等题材。",
        "支持已批准制片厂提交视频，审片员或管理员审核，审核通过后发布。",
        "普通内容查询只返回 PUBLISHED 状态，支持按类型、题材和标题关键字筛选。",
        "密码写入时使用 BCrypt 加密，Swagger 接口文档可访问。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.3 当前关键缺口", 2)
    add_table(
        doc,
        ["缺口", "实际影响", "严重度"],
        [
            ("无登录与 JWT 实现", "前端已声明 /auth/login，但后端没有对应接口；JwtUtil 为空。", "阻断"),
            ("所有接口 permitAll", "任何人都可修改角色、审核制片厂、审核和发布视频。", "阻断"),
            ("前端无法构建", "App.tsx 缺失，main.tsx 未渲染应用，api.ts 存在 TS 错误。", "阻断"),
            ("无真实媒体链路", "仅保存 URL；没有上传、对象存储、转码、切片、清晰度和播放鉴权。", "高"),
            ("评论模块为空", "实体、仓库、服务、控制器均无实现。", "中"),
            ("AI 服务不存在", "推荐、语义搜索、字幕生成目前只是 README 目标。", "中"),
            ("配置含明文密码", "数据库密码和 JWT 示例密钥写入配置，生产环境存在泄露风险。", "高"),
            ("异常与审计不足", "全局异常处理为空，审核操作没有不可抵赖的审计日志。", "高"),
        ],
        [1.45, 4.1, 0.95],
        header_fill=LIGHT_RED,
    )

    add_heading(doc, "3. 项目优点与不足", 1)
    add_heading(doc, "3.1 主要优点", 2)
    for item in [
        "角色和业务边界明确。制片厂负责内容供给，审片员负责审核，管理员负责治理，普通用户负责消费，符合平台型产品的组织模型。",
        "审核状态机方向正确。视频包含草稿、待审、通过、拒绝、发布和下架状态，为后续补充撤回、复审和版本管理留下空间。",
        "技术栈主流且学习资源丰富。Spring Boot、React、PostgreSQL、Redis 和 Docker 适合中小规模项目迭代。",
        "实体字段已覆盖基础运营指标。视频已有播放量、点赞量、评论量、评分、审核人和发布时间等字段。",
        "开发基础设施较完整。数据库、缓存和管理工具均可通过 Compose 启动，降低本地环境搭建成本。",
        "后端分层清晰。Controller、Service、Repository、DTO、Entity 的结构利于继续扩展。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.2 主要不足", 2)
    for item in [
        "当前安全性是最大短板：业务层虽然检查角色，但请求中的 userId、reviewerId 可由客户端伪造，无法代表真实登录身份。",
        "数据模型仍偏“单视频文件”，没有电视剧季/集、综艺期数、演职员、地区、语言、版权窗口、付费模式等流媒体核心模型。",
        "视频搜索仍是数据库标题模糊查询，不支持分页、排序、标签、多字段检索和中文全文搜索。",
        "上传限制为 100MB，不适合真实长视频；同时没有对象存储与异步转码任务。",
        "前后端契约出现漂移：前端声明认证接口，后端未实现；README 的端口、AI 服务和功能列表与仓库事实不一致。",
        "工程质量不足：中文编码乱码、空类、无迁移工具、无测试覆盖、生产日志级别过高。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 产品目标与范围", 1)
    add_heading(doc, "4.1 产品目标", 2)
    add_number(doc, "建立受控的内容供应链，使已批准制片厂可以提交片库并追踪审核结果。")
    add_number(doc, "建立可审计的内容审核流程，使审片员只能处理授权范围内的待审内容。")
    add_number(doc, "建立面向普通用户的独立观影站点，只展示可播放且处于版权有效期内的内容。")
    add_number(doc, "建立管理员治理能力，覆盖用户、角色、制片厂、内容、分类、下架和审计。")
    add_number(doc, "在核心业务稳定后，引入 AI 搜索、推荐和字幕能力，而不是让 AI 成为首期阻塞项。")

    add_heading(doc, "4.2 MVP 范围", 2)
    add_table(
        doc,
        ["纳入 MVP", "暂不纳入 MVP"],
        [
            ("注册、登录、JWT、角色权限", "会员订阅、付费点播、广告结算"),
            ("制片厂申请与管理员审批", "复杂合同、分账和发票管理"),
            ("视频元数据、封面和媒体上传", "直播、弹幕和多人互动"),
            ("审核、退回、发布、下架", "多区域 CDN 智能调度"),
            ("片库搜索、筛选、详情、播放", "高级 AI 推荐和大模型问答"),
            ("播放记录、收藏、基础点赞", "社交关系链和创作者社区"),
            ("操作日志和基础监控", "复杂 BI 数据仓库"),
        ],
        [3.25, 3.25],
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "5. 用户角色与权限需求", 1)
    add_table(
        doc,
        ["角色", "核心职责", "允许操作", "禁止或受限操作"],
        [
            ("管理员", "平台治理与授权", "管理用户角色、审批制片厂、查看全部内容、发布/下架、查看审计", "不得绕过审计记录"),
            ("审片员", "内容合规审核", "查看待审队列、播放样片、通过/拒绝、填写意见", "不能修改用户角色和制片厂资质"),
            ("制片厂", "内容供应与维护", "申请资质、提交片库、查看自有内容和审核结果、修订重提", "不能审核或查看其他制片厂私有内容"),
            ("普通用户", "内容消费", "注册登录、搜索、观看、收藏、点赞、查看历史", "不能访问待审、拒绝或下架内容"),
        ],
        [0.85, 1.35, 2.35, 1.95],
        font_size=8.8,
    )

    add_callout(
        doc,
        "权限实现原则",
        "所有身份信息必须从 JWT 或服务端会话获得，不能信任请求体中的 createdBy、reviewerId 或任意 userId。"
        "接口层使用角色授权，服务层继续校验资源归属和状态流转，关键操作写入审计日志。",
        LIGHT_GREEN,
    )

    add_heading(doc, "6. 核心业务流程", 1)
    add_heading(doc, "6.1 制片厂准入流程", 2)
    for step in [
        "普通用户填写公司名称、简介、联系人和证明材料，提交制片厂申请。",
        "系统将申请状态置为 PENDING，禁止重复提交未完成申请。",
        "管理员查看材料并选择通过或拒绝；拒绝必须填写原因。",
        "通过后用户角色转为 STUDIO，状态为 APPROVED，并获得片库提交权限。",
        "所有审批动作记录操作者、时间、申请版本、意见和结果。",
    ]:
        add_number(doc, step)

    add_heading(doc, "6.2 视频审核发布流程", 2)
    for step in [
        "已批准制片厂创建内容条目，填写类型、题材、简介、版权信息、封面和媒体文件。",
        "媒体上传后进入异步处理：病毒扫描、元数据提取、转码、截图和播放可用性检查。",
        "制片厂提交审核后状态变为 PENDING，提交版本在审核完成前不可直接覆盖。",
        "审片员观看样片、核对材料，选择 APPROVED 或 REJECTED，并填写审核意见。",
        "审核通过后由制片厂或管理员执行发布；发布时校验媒体处理完成和版权窗口有效。",
        "管理员可因版权、违规或运营原因下架，且必须记录原因。",
    ]:
        add_number(doc, step)

    add_heading(doc, "6.3 普通用户观看流程", 2)
    for step in [
        "用户访问独立观影站，可匿名浏览公开片库；播放和个性化功能要求登录。",
        "搜索支持关键字、类型、题材、地区、年份和排序，并提供分页。",
        "详情页展示简介、演职员、剧集列表、评分和播放入口。",
        "播放服务校验内容状态、版权窗口、用户权限和媒体可用性，再签发短期播放地址。",
        "系统异步记录播放进度、观看历史和基础行为事件。",
    ]:
        add_number(doc, step)

    add_heading(doc, "7. 功能需求清单", 1)
    functional_rows = [
        ("FR-01", "账户注册与登录", "支持用户名/邮箱注册、BCrypt 密码、JWT 登录和退出。", "P0", "未实现"),
        ("FR-02", "角色授权", "按 ADMIN/REVIEWER/STUDIO/USER 限制接口和页面。", "P0", "未实现"),
        ("FR-03", "用户管理", "管理员查询、启停用户、修改角色、重置状态。", "P0", "部分实现"),
        ("FR-04", "制片厂申请", "用户提交资质，管理员审核并记录原因。", "P0", "部分实现"),
        ("FR-05", "内容建档", "制片厂维护片名、类型、题材、简介、版权和主创信息。", "P0", "部分实现"),
        ("FR-06", "媒体上传处理", "对象存储上传、异步转码、封面和播放文件生成。", "P0", "未实现"),
        ("FR-07", "审核工作台", "待审队列、样片播放、审核意见、通过/拒绝。", "P0", "部分实现"),
        ("FR-08", "发布与下架", "仅通过内容可发布；管理员可下架并保留原因。", "P0", "部分实现"),
        ("FR-09", "公开片库", "只展示可播放的 PUBLISHED 内容，支持分页筛选。", "P0", "部分实现"),
        ("FR-10", "视频播放", "签名 URL/HLS 播放、断点续播、清晰度切换。", "P0", "未实现"),
        ("FR-11", "收藏与历史", "用户收藏、取消收藏、记录观看进度和历史。", "P1", "未实现"),
        ("FR-12", "评论与评分", "发布评论、审核/删除、评分聚合。", "P1", "未实现"),
        ("FR-13", "运营统计", "用户、内容、审核时效、播放趋势等基础指标。", "P1", "未实现"),
        ("FR-14", "AI 语义搜索", "基于标题、简介、字幕向量进行语义召回。", "P2", "未实现"),
        ("FR-15", "AI 推荐", "基于行为与内容特征生成推荐列表，并提供降级策略。", "P2", "未实现"),
        ("FR-16", "AI 字幕", "语音转写、时间轴字幕、人工修订和发布。", "P2", "未实现"),
    ]
    add_table(
        doc,
        ["编号", "模块", "需求说明", "优先级", "现状"],
        functional_rows,
        [0.62, 1.18, 3.45, 0.58, 0.67],
        header_fill=LIGHT_BLUE,
        font_size=8.25,
    )

    add_heading(doc, "8. 视频领域模型需求", 1)
    add_heading(doc, "8.1 分类模型", 2)
    add_body(
        doc,
        "现有 VideoType 与 VideoGenre 的拆分是正确方向，但应避免将所有内容都压在单张 videos 表。"
        "建议将“作品”和“可播放媒体”分离，以支持电视剧多季多集、综艺多期和同一作品多清晰度。"
    )
    add_table(
        doc,
        ["实体", "核心字段", "说明"],
        [
            ("Content", "标题、类型、简介、年份、地区、语言、状态、制片厂", "作品级信息"),
            ("Season", "作品、季号、名称、排序", "电视剧和综艺可选"),
            ("Episode", "季、集号、标题、时长、发布时间", "单集内容"),
            ("MediaAsset", "原文件、HLS 地址、清晰度、码率、处理状态", "可播放媒体"),
            ("Genre/Tag", "名称、状态、排序", "题材与运营标签"),
            ("Copyright", "授权区域、开始/结束时间、证明材料", "发布前强校验"),
            ("ReviewRecord", "版本、审核人、结果、意见、时间", "不可覆盖的审核历史"),
        ],
        [1.15, 3.2, 2.15],
        font_size=8.8,
    )

    add_heading(doc, "8.2 状态机约束", 2)
    add_body(
        doc,
        "建议状态流转为 DRAFT → PROCESSING → PENDING → APPROVED/REJECTED → PUBLISHED → BANNED/EXPIRED。"
        "拒绝内容修改后应形成新版本并重新提交，而不是直接覆盖原审核证据。"
    )

    add_heading(doc, "9. 非功能需求", 1)
    add_table(
        doc,
        ["类别", "最低要求", "验收口径"],
        [
            ("安全", "JWT、RBAC、资源归属校验、密码加密、限流、审计日志", "越权接口测试全部失败，关键操作可追溯"),
            ("性能", "列表接口分页；常用查询建立索引；异步处理媒体", "普通查询 P95 小于 500ms，不含 AI 和转码"),
            ("可用性", "服务健康检查、明确错误码、上传可重试", "依赖故障时返回可理解错误，不产生脏状态"),
            ("可维护性", "数据库迁移、统一异常、OpenAPI、模块化前端", "新环境可重复部署，接口契约可生成"),
            ("可观测性", "结构化日志、指标、追踪 ID、错误告警", "一次请求可跨日志定位，审核失败可回溯"),
            ("数据保护", "密钥环境变量化、最小权限、备份恢复", "仓库不含真实密码，恢复演练可执行"),
            ("兼容性", "桌面与移动端响应式，主流浏览器", "核心流程在 Chrome/Edge 和移动视口可用"),
        ],
        [1.05, 3.15, 2.3],
        font_size=8.8,
    )

    add_heading(doc, "10. AI 微服务需求边界", 1)
    add_body(
        doc,
        "AI 微服务应是可替换的增强服务，而不是账户、审核和播放主链路的单点依赖。"
        "Java 主服务维护权威业务数据，Python 服务负责模型调用、向量检索和异步任务。"
    )
    add_table(
        doc,
        ["能力", "输入", "输出", "降级策略"],
        [
            ("语义搜索", "标题、简介、字幕、用户查询", "候选内容 ID 与相关度", "退回数据库关键字搜索"),
            ("内容推荐", "用户行为、内容标签、热度", "排序后的内容 ID", "热门与分类推荐"),
            ("字幕生成", "音频或视频文件", "带时间轴字幕文件", "任务失败后人工上传字幕"),
            ("内容辅助审核", "文本、封面、抽帧", "风险标签与解释", "仅作提示，不自动替代审片员决定"),
        ],
        [1.15, 2.0, 1.75, 1.6],
        font_size=8.8,
    )
    add_callout(
        doc,
        "AI 合规要求",
        "模型输出必须记录模型版本、请求时间和置信度；审核建议必须由人工最终确认；"
        "不得将未授权视频或个人信息直接发送给第三方模型服务。",
        LIGHT_YELLOW,
    )

    add_heading(doc, "11. 接口与前端需求", 1)
    add_heading(doc, "11.1 接口原则", 2)
    for item in [
        "统一返回结构、业务错误码和 HTTP 状态码；禁止只抛 IllegalArgumentException。",
        "所有列表接口支持 page、size、sort，并限制最大分页大小。",
        "当前用户身份通过认证上下文获得，不从请求体接收 createdBy 或 reviewerId。",
        "上传采用预签名地址或分片上传，避免大文件经过 Java 应用内存。",
        "OpenAPI 文档与前端 TypeScript 类型自动生成，减少契约漂移。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11.2 前端站点划分", 2)
    add_table(
        doc,
        ["站点", "目标用户", "核心页面"],
        [
            ("普通用户站", "普通用户/游客", "首页、搜索、分类、详情、播放、收藏、历史、个人中心"),
            ("业务工作台", "制片厂/审片员", "登录、制片厂申请、片库管理、上传、审核队列、审核详情"),
            ("管理后台", "管理员", "仪表盘、用户权限、制片厂审批、内容治理、分类、审计与配置"),
        ],
        [1.25, 1.45, 3.8],
    )

    add_heading(doc, "12. 安全与合规风险", 1)
    add_table(
        doc,
        ["风险", "当前表现", "整改要求", "优先级"],
        [
            ("水平越权", "接口使用客户端传入的用户 ID", "从认证上下文取身份并校验资源归属", "P0"),
            ("垂直越权", "SecurityConfig 对所有请求 permitAll", "按角色配置 URL 与方法级权限", "P0"),
            ("凭据泄露", "数据库密码和 JWT 密钥写在 YAML/Compose", "使用环境变量和本地 .env，生产用密钥服务", "P0"),
            ("媒体盗链", "直接保存公开 URL", "签名 URL、短期令牌、域名和来源限制", "P1"),
            ("版权过期", "无授权区域与有效期模型", "发布和播放时校验版权窗口", "P1"),
            ("审核不可追溯", "Video 只保存最后一次审核信息", "独立 ReviewRecord 和操作日志", "P1"),
            ("数据迁移风险", "Hibernate ddl-auto=update", "引入 Flyway/Liquibase 并版本化迁移", "P1"),
        ],
        [1.05, 1.75, 2.8, 0.9],
        font_size=8.5,
        header_fill=LIGHT_RED,
    )

    add_heading(doc, "13. 迭代路线图", 1)
    add_table(
        doc,
        ["阶段", "周期建议", "交付目标", "退出条件"],
        [
            ("阶段 0：修复基线", "1 周", "前端可构建、配置去敏、统一异常、补数据库迁移", "CI 编译通过，开发环境一键启动"),
            ("阶段 1：身份与权限", "1-2 周", "注册登录、JWT、RBAC、当前用户接口、审计日志", "越权测试通过"),
            ("阶段 2：内容供应链", "2-3 周", "制片厂申请、内容建档、上传、转码、审核、发布", "完整流程可演示并可回溯"),
            ("阶段 3：观影 MVP", "2 周", "搜索、详情、HLS 播放、历史、收藏", "普通用户可稳定观看已发布内容"),
            ("阶段 4：运营能力", "1-2 周", "评论、评分、统计、下架、监控告警", "后台可完成日常治理"),
            ("阶段 5：AI 增强", "2-4 周", "语义搜索、推荐、字幕任务与降级", "AI 故障不影响主业务"),
        ],
        [1.25, 0.85, 2.9, 1.5],
        font_size=8.6,
        header_fill=LIGHT_BLUE,
    )

    add_heading(doc, "14. MVP 验收标准", 1)
    for item in [
        "未登录用户不能访问管理、审核和制片厂私有接口。",
        "普通用户无法通过修改请求参数冒充制片厂或审片员。",
        "未批准制片厂不能提交视频，已批准制片厂只能管理自有内容。",
        "只有待审内容可被审核，只有审核通过且媒体处理完成的内容可发布。",
        "普通用户站点无法查询或播放草稿、待审、拒绝、下架和版权过期内容。",
        "视频上传失败、转码失败和审核拒绝均有明确状态与可恢复操作。",
        "所有管理员和审片员关键操作都有操作者、时间、对象、前后状态和原因。",
        "前后端均能在干净环境构建，Docker 开发环境可重复启动。",
        "核心接口至少覆盖服务层单元测试和控制器权限测试。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "15. 结论与建议", 1)
    add_body(
        doc,
        "该项目的业务方向、角色划分和审核流程具备良好基础，适合作为视频平台 MVP 的起点。"
        "当前最值得保留的是用户角色模型、制片厂准入思路、视频类型/题材拆分以及审核状态流。"
    )
    add_body(
        doc,
        "下一步不应立即投入复杂 AI 推荐，而应优先解决认证授权、前端可构建、媒体上传播放、审核审计和数据库迁移。"
        "这些能力决定平台是否真实可用，也决定后续 AI 是否有可靠数据和业务接口可依赖。"
    )
    add_callout(
        doc,
        "推荐决策",
        "将项目目标从“完整 AI 流媒体平台”收敛为“具备内容准入、审核、发布和播放闭环的流媒体 MVP”。"
        "先用 6 至 8 周形成可验收主链路，再将 AI 作为独立增强阶段推进。",
        LIGHT_GREEN,
    )

    add_heading(doc, "附录 A：分析依据", 1)
    for item in [
        "后端 28 个 Java 源文件及 Maven 配置。",
        "User、Video、Role、StudioStatus、VideoStatus、VideoType、VideoGenre 等领域模型。",
        "UserController、VideoController、UserService、VideoService 的实际业务逻辑。",
        "SecurityConfig、JwtUtil、异常处理、评论模块的当前实现状态。",
        "React/TypeScript 依赖、main.tsx、services/api.ts 与实际构建结果。",
        "application.yml、开发/生产 Profile、Dockerfile 与 Docker Compose 配置。",
        "PostgreSQL、Redis、pgAdmin、Redis Commander 的当前容器运行状态。",
        "README 中的目标描述与当前仓库实现之间的差异。",
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = "视频流媒体平台需求分析报告"
    doc.core_properties.subject = "Video AI Platform 现状审计、需求范围与迭代规划"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "视频平台, 需求分析, Spring Boot, React, AI"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
